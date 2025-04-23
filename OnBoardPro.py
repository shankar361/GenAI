from docx import Document as DocxDocument
import streamlit as st
from langchain_openai import ChatOpenAI
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.chains.retrieval_qa.base import RetrievalQA
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyMuPDFLoader
from pdf2docx import Converter
from langchain.schema import Document
import tempfile
from datetime import datetime
import dotenv
import os

dotenv.load_dotenv()
api_key   = os.getenv("api_key")
if "intro" not in st.session_state:
    st.session_state.intro=True
if "chat_history" not in st.session_state:
    st.session_state.chat_history = ""

if "singleConv" not in st.session_state:
    st.session_state.singleConv = []

if "query_input" not in st.session_state:
    st.session_state.query_input = ""

if "prevQuery" not in st.session_state:
    st.session_state.prevQuery=""

if "query" not in st.session_state:
    st.session_state.query = ""

if "chat_query" not in st.session_state:
    st.session_state.chat_query=""
if "docs" not in st.session_state:
    st.session_state.docs = []

if "qa" not in st.session_state:
    st.session_state.qa = None
if "uploaded_files" not in st.session_state:
    st.session_state.uploaded_files = []

if "resumeUploaded" not in st.session_state:
    st.session_state.resumeUploaded=False

if "resume_verified" not in st.session_state:
     st.session_state.resume_verified = False
if "clear_chat_clicked" not in st.session_state:
     st.session_state.clear_chat_clicked = False 

if not api_key:
     st.warning("API key not found , please check environment variables")
else:    
    if "llm" not in st.session_state:    
        st.session_state.llm = ChatOpenAI(
        base_url  = "https://api.groq.com/openai/v1",
        api_key   = api_key,
      # model_name= "llama3-70b-8192"  #model name may reach limit
    #    model_name= "llama3-8b-8192"  # reached limit
        model_name="llama-3.1-8b-instant"
     #  model_name= "llama3-70b-8192"  #reached limit
    )   
        
def clear_text():
    #before clearing the textbox, store the input to "query"
    st.session_state.query = st.session_state.query_input
    st.session_state.query_input = ""

def buildResult(response):
                    st.write("🧑You : "+response["query"])   
                    st.write("🤖HRBuddy : "+response["result"]) 
                    q =   response["query"]
                    r = response["result"]
                    #put/append into chat box                                                              
                    line = '_' * 80                    
                    # st.text_area("Conversation with HRBuddy", value=st.session_state.chat_history, height=400)
                    st.write("Full Conversation with HRBuddy (Download from left panel):")
                    r +="\n"+ line
                    st.session_state.singleConv.append((q,r))           
                    render_chat(st.session_state.singleConv)

def CheckIfResumeUploaded(qa):
 if st.session_state.resume_verified == False:
        resp = qa.invoke("Is this a resume/CV, answer in yes or no.")
        ans = resp["result"].lower()
        st.session_state.resume_verified = True
        if "yes" not in ans:
            st.error("The candidate's resume is not uploaded! Please upload resume/CV")         
            st.session_state.resume_verified = False 
            return False
        else:
            return True
        
def start_chat(qa,resumeUploaded):
    if  st.session_state.clear_chat_clicked :
         return
   
    if resumeUploaded:
        query = st.session_state.query
       # query = st.session_state.query_input.strip()
        upQuery = query.upper()
        if upQuery == "Q" or upQuery == "QUIT" or upQuery == "END" or upQuery == "BYE" :
            st.write("Thank you for using our AI system, Good bye..!")
        else:
            if query != "" :
                if query != st.session_state.prevQuery:
                    with st.spinner("Analyzing data based on your query.."):
                        response = qa.invoke(query)
                        #if not response["source_documents"]:
                            #st.write("The Question is out of context or the correct documents not provided")
                        #else:
                        buildResult(response)
                        st.session_state.prevQuery = query
                    #  st.session_state.query = ""
                else:
                    render_chat(st.session_state.singleConv)
                  
            

def render_chat(singleConv: str):
    st.markdown("""
    <style>
    .chat-box {
        max-height: 400px;
        overflow-y: auto;
        color     : #784212;
        border    : 1px solid #444;
        border-radius: 5px;    
    }
    .chat-message-user {
    color: #1f77b4;
    margin-bottom: 8px;
    }
    .chat-message-bot {
    color: #2ca02c;
    margin-bottom: 12px;
    }
    </style>
    """, unsafe_allow_html=True)

    chat_html = "<div class='chat-box'>"
    for query,response in singleConv:
        chat_html += f"<div class='chat-message-user'>🧑 You: {query}</div>"
        chat_html += f"<div class='chat-message-bot'>🤖 HRBuddy: {response}</div>"
    chat_html += "</div>"

    st.markdown(chat_html, unsafe_allow_html=True)



def load_file(file_path):
    docx  = DocxDocument(file_path)
   
    text  = "\n".join([p.text for p in docx.paragraphs if p.text.strip()])
    tables_text = extractTableData(file_path)
    fullText = text +"\n"+ tables_text  
    return [Document(page_content=fullText)]

def splitData(pdfData):
   splitter = RecursiveCharacterTextSplitter(chunk_size=500,chunk_overlap=100)
   splittedpdfDoc = splitter.split_documents(pdfData)
   return splittedpdfDoc

def extractTableData(file_path):
    tables_text=[]

    docx2=DocxDocument(file_path)
    for tables in docx2.tables:
        for rows in tables.rows:
           # for cells in rows.cells:
            row_data = [cell.text.strip() for cell in rows.cells]
            tables_text.append(" | ".join(row_data))
        tables_text.append("-"*50)

    return "\n".join(tables_text)


def loadAndSplit_pdfFile(files): 
    loader  = PyMuPDFLoader(files)
    pdfData = loader.load()
    splittedpdfDoc = splitData(pdfData)
    return splittedpdfDoc

def createTempfile(files):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        tmp_file.write(files.read())
        return tmp_file.name

def  createRetrieverQA(docs,emb):
    vectorStores = FAISS.from_documents(docs,emb)
    vectorStores.save_local("faiss")
    retriever = vectorStores.as_retriever()
    qa = RetrievalQA.from_llm(
    llm = st.session_state.llm,
    retriever = retriever,
    return_source_documents=False       
    )
    return qa


if "emb" not in st.session_state:
    st.session_state.emb= HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

st.title("OnBoarding Pro")
st.write("Dear HR, How can I help you today?")
introText = "This is an HR assistance tool which works based on Candidates' profile" \
"  and interview feedbacs. Please uploaded all supporting document to have best results"

#st.header("******Your Personalized Hiring assistant******")
path = st.sidebar.file_uploader("Upload Candidate's Resume and supporting docs(*.pdf, *.docx)",
accept_multiple_files = True,
help = "Upload one more more files"
)
if st.session_state.intro and path ==[]:
    st.write(introText)
    st.session_state.intro = False

st.text_input("Enter your query here", key="query_input", on_change=clear_text)

if st.sidebar.button("Clear Chat"):
    st.text_area("Conversation with HRBuddy", value="", height=400)
    st.session_state.clear_chat_clicked = True
    st.session_state.singleConv = []                     
else:          
    st.session_state.clear_chat_clicked = False
    
#st.write(st.session_state.chat_history)  #debug
if  st.session_state.singleConv != []:
    for q,r in st.session_state.singleConv:
        textBytes = q.encode('utf-8')
        textBytes+= r.encode('utf-8')
  #  st.write(textBytes)
    now1 = datetime.now()
    st.sidebar.download_button(label = "Download Chat",
    data      = textBytes,
    file_name = "downloadedChat_"+ now1.strftime("%Y-%m-%d-%H:%M:%S")+".txt",
    mime      = "text/plain"
    )
    

new_docs = []
if path is not None:  
    for files in path:
        tmp = createTempfile(files)
        if files.name.endswith(".pdf"):     
           
#convert pdf to docx because model works best with docx
            docxTemp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            docxPath = docxTemp.name
            cv = Converter(tmp)
            cv.convert(docxPath,0,end=None)   
        
          #  splittedpdfDoc = loadAndSplit_pdfFile(tmp2)
            new_docs.extend(load_file(docxPath))
        else:
            new_docs.extend(load_file(tmp))

files_uploaded = [f.name for f in path]if path else []
files_changed  = (files_uploaded!= st.session_state.uploaded_files)

if len(new_docs) > 0 and files_changed:
   # st.session_state.query = ""
    st.session_state.uploaded_files=files_uploaded
    st.session_state.docs = new_docs
    st.session_state.resume_verified = False   
    with st.spinner("Building dataset, please wait..."):       
        st.session_state.qa =  createRetrieverQA(new_docs,st.session_state.emb) 
        st.session_state.resumeUploaded = CheckIfResumeUploaded(st.session_state.qa)  

if st.session_state.qa: 
    start_chat(st.session_state.qa,st.session_state.resumeUploaded)
   
if not path:
     st.info("Upload a resume or other supporting documents to begin")
