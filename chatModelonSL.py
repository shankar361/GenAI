from docx import Document as DocxDocument
import speech_recognition as SR
import streamlit as st
from langchain_openai import ChatOpenAI
from langchain_community.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings
from sentence_transformers import SentenceTransformer
from langchain.chains.retrieval_qa.base import RetrievalQA
#from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyMuPDFLoader
from pdf2docx import Converter
from langchain.schema import Document
import tempfile
from datetime import datetime
import dotenv
import os

dotenv.load_dotenv()
api_key   = os.getenv("api_key")

if "intro" not in st.session_state:
    st.session_state.intro=True

if "chat_history" not in st.session_state:
    st.session_state.chat_history = ""

if "singleConv" not in st.session_state:
    st.session_state.singleConv = []

if "query_input" not in st.session_state:
    st.session_state.query_input = ""

if "prevQuery" not in st.session_state:
    st.session_state.prevQuery=""

if "query" not in st.session_state:
    st.session_state.query = ""

if "chat_query" not in st.session_state:
    st.session_state.chat_query=""
if "docs" not in st.session_state:
    st.session_state.docs = []

if "qa" not in st.session_state:
    st.session_state.qa = None

if "uploaded_files" not in st.session_state:
    st.session_state.uploaded_files = []

if "firstPromptGiven" not in st.session_state:
    st.session_state.firstPromptGiven=False

if "resumeUploaded" not in st.session_state:
    st.session_state.resumeUploaded=False

if "resume_verified" not in st.session_state:
     st.session_state.resume_verified = False

if "clear_chat_clicked" not in st.session_state:
     st.session_state.clear_chat_clicked = False 

if not api_key:
     st.error("API key not found , please check environment variables")
     st.stop()
else:    
    if "llm" not in st.session_state:    
        st.session_state.llm = ChatOpenAI(
        base_url  = "https://api.groq.com/openai/v1",
        api_key   = api_key,
      # model_name= "llama3-70b-8192"  #model name may reach limit
    #    model_name= "llama3-8b-8192"  # reached limit
      #  model_name = st.sidebar.selectbox("Choose a LLM model",["llama3-8b-8192","llama-3.1-8b-instant"])
        model_name="llama-3.1-8b-instant"
     #  model_name= "llama3-70b-8192"  #reached limit
    )   
        
def clear_text():
    #before clearing the textbox, store the input to "query"
    st.session_state.query = st.session_state.query_input
    st.session_state.query_input = ""

def setQuery():
    st.session_state.query = st.session_state.promptSelect

def buildResult(response):
    st.write("🧑 You : "+response["query"])   
    st.write("🧠 HRBuddy : "+response["result"]) 
    q = "\n"
    q += response["query"]
    r  = response["result"]                                                       
    line = '_' * 80                   
    # st.text_area("Conversation with HRBuddy", value=st.session_state.chat_history, height=400)
    if not st.session_state.clear_chat_clicked:
        with st.expander("Full Conversation with HRBuddy (Download from left panel)"):
        # st.write("Full Conversation with HRBuddy (Download from left panel):")
            r +="\n"+ line
            st.session_state.singleConv.append((q,r))           
            render_chatST(st.session_state.singleConv)

def CheckIfResumeUploaded(qa):
 if st.session_state.resume_verified == False:
    try:

        resp = qa.invoke(
        """Analyze the uploaded documents. Based on its structure and content,
        is this likely a resume or CV of a person?
        Respond with only 'Yes' or 'No'."""
    )
    except Exception as e:
        resp=""
        if "please reduce your message size" in str(e).lower or "request too large for" in str(e).lower:
            st.error("Request too large for the model, upgrade your tier!")
        else:
            st.error("Error occureed while sending query! "+str(e))
    if not resp:
        return
    ans = resp["result"].lower()
    st.session_state.resume_verified = True
    if "yes" not in ans:
        st.error("The candidate's resume is not uploaded! Please upload resume/CV")         
        st.session_state.resume_verified = False 
        return False
    else:
        return True

def suggestPrompts2(qa):
    st.session_state.firstPromptGiven = True
    suggPrompt="Suggest 3-5 short prompts based on provided documents,focus only on selection, rejection status and interview feedbacks,Dont give any reason" 
    try:
        prompts = qa.invoke(suggPrompt)   
        newlines=[] 
        lines = prompts["result"]
     #   st.write(lines)
        singleLines=lines.splitlines()
        for eachLine in  singleLines:
            if eachLine !="":
                newlines.append(eachLine)
        st.selectbox("Select an AI generated prompt or type a query",newlines[0:],key="promptSelect",on_change=setQuery)  
      
    except:
        st.error("No prompts available, ener a query to proceed")

def start_chat(qa,resumeUploaded):  
    if  st.session_state.clear_chat_clicked :
         return  
    if resumeUploaded:
       # lines = ""
        query = st.session_state.query
        if not st.session_state.firstPromptGiven:
            suggestPrompts2(qa) 
        if query:
            query = query[0].upper() + query[1:]
       # query = st.session_state.query_input.strip()
        upQuery = query.upper()
        if upQuery == "Q" or upQuery == "QUIT" or upQuery == "END" or upQuery == "BYE" :
            st.write("Thank you for using our AI system, Good bye..!")
        else:
            if query != "" :
                if query != st.session_state.prevQuery:
                    with st.spinner("Analyzing data based on your query.."):
                        try:
                            response = qa.invoke(query)
                        except Exception as e:
                            if "please reduce your message size" in str(e).lower or "request too large for" in str(e).lower:
                                st.error("Request too large for the model, upgrade your tier!")
                            else:
                                st.error("Error occureed while sending query! "+str(e))

                       # st.write("Source docs:", response["source_documents"])
                        #if not response["source_documents"]:
                        #st.write("The Question is out of context or the correct documents not provided")
                        #else:
                        buildResult(response)
                        st.session_state.prevQuery = query
                    #  st.session_state.query = ""
                else:
                    if not st.session_state.clear_chat_clicked:
                        with st.expander("Full Conversation with HRBuddy (Download from left panel)"):
                            render_chatST(st.session_state.singleConv)
                  
            
def render_chatST(singleConv: str):
    for query,response in singleConv:
        with st.chat_message("user",):
            st.markdown(query)
        with st.chat_message("ai"):
            st.markdown(response)

def render_chat(singleConv: str):
    st.markdown("""
    <style>
    .chat-box {
        max-height: 400px;
        overflow-y: auto;
        color     : #784212;
        border    : 1px solid #444;
        border-radius: 5px;    
    }
    .chat-message-user {
    color: #1f77b4;
    margin-bottom: 8px;
    }
    .chat-message-bot {
    color: #2ca02c;
    margin-bottom: 12px;
    }
    </style>
    """, unsafe_allow_html = True)

    chat_html = "<div class='chat-box'>"
    for query,response in singleConv:
        chat_html += f"<div class='chat-message-user'>🧑  You: {query}</div>"
        chat_html += f"<div class='chat-message-bot'>🤖  HRBuddy: {response}</div>"
    chat_html += "</div>"
    st.markdown(chat_html, unsafe_allow_html = True)



def load_file(file_path,source_name):
    docx  = DocxDocument(file_path)
   
    text  = "\n".join([p.text for p in docx.paragraphs if p.text.strip()])
    tables_text = extractTableData(file_path)
    fullText = text +"\n"+ tables_text  
    return [Document(page_content=fullText,metadata={"source": source_name})]

#def splitData(pdfData):
#   splitter = RecursiveCharacterTextSplitter(chunk_size = 1000,chunk_overlap=150)
#   splittedpdfDoc = splitter.split_documents(pdfData)
#   return splittedpdfDoc

def extractTableData(file_path):
    tables_text = []

    docx2=DocxDocument(file_path)
    for tables in docx2.tables:
        for rows in tables.rows:
           # for cells in rows.cells:
            row_data = [cell.text.strip() for cell in rows.cells]
            tables_text.append(" | ".join(row_data))
        tables_text.append("-"*50)

    return "\n".join(tables_text)


def loadAndSplit_pdfFile(files): 
    loader  = PyMuPDFLoader(files)
    pdfData = loader.load()
  #  splittedpdfDoc = splitData(pdfData)
    return pdfData

def createTempfile(files):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        tmp_file.write(files.read())
        return tmp_file.name
    

def inputFromMicrophone():
    recog = SR.Recognizer()
    try:
        with SR.Microphone() as source:
            with st.spinner("Listening..."):
                recog.adjust_for_ambient_noise(source,duration=1)
                audio = recog.listen(source,timeout = None,phrase_time_limit = 20)
  
        st.session_state.query = recog.recognize_google(audio)
    except AttributeError as e:
        st.error("Could not find PyAudio; check audio hardware installation")
    except SR.UnknownValueError as e:
        st.warning("Coud not understand audio, use textbox instead or try again!")
    except SR.RequestError as e:    
        st.error("Could not get response from Google Speech Recognistion!")
    except:
        st.error("Other errors!")

def  createRetrieverQA(docs,emb):
    vectorStores = FAISS.from_documents(docs,emb)
  #  vectorStores.save_local("faiss")
    retriever = vectorStores.as_retriever()
    try:

        qa = RetrievalQA.from_llm(
    llm = st.session_state.llm,
    retriever = retriever,
    return_source_documents=True       
    )
    except Exception as e:
        st.error("Error occured while creating RetrievalQA "+e)
    return qa


if "emb" not in st.session_state:
    st.session_state.emb= HuggingFaceEmbeddings(model_name="sentence-transformers/all-mpnet-base-v2")
   
 #   model = SentenceTransformer("sentence-transformers/all-mpnet-base-v2")
   # st.session_state.emb = HuggingFaceEmbeddings(model_name=model)
st.title("OnBoarding Pro")
st.subheader("Dear HR, How can I help you today?")
introText = "This is an HR assistance tool which works based on Candidates' profile" \
"  and interview feedbacks. Please upload all supporting document to have best results"

with st.sidebar:
    colImg1,colImg2,colImg13=st.columns(3)
    with colImg2:
        st.image("silicon.png")
    st.markdown("---")
path = st.sidebar.file_uploader("Upload Candidate's Resume and supporting docs(*.pdf, *.docx)",
accept_multiple_files = True,
help = "Upload one more more files"
)
if st.session_state.intro and path ==[]:
    st.write(introText)
    st.session_state.intro = False
col1, col2 = st.columns([5, 1])
with col1:
    st.text_input("💡Enter your query here", 
                  key="query_input", on_change=clear_text,
                  placeholder="Eg. Is the candidate selected?")
with col2:
    st.write("")
    if st.button("🎤",help = "Click here to speak your query"): 
        inputFromMicrophone()
if st.session_state.qa: 
    if st.button("💡Suggest Prompts",help = "Feeling lazy? Click here to show AI generated prompts"):
       # suggestPrompts(st.session_state.qa)
        st.session_state.firstPromptGiven=False

new_docs = []
if path is not None:  
    for files in path:
        tmp = createTempfile(files)
        if files.name.endswith(".pdf"):     
           
            docxTemp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            docxPath = docxTemp.name
            cv = Converter(tmp)
            cv.convert(docxPath,0,end=None)   
        
          #  splittedpdfDoc = loadAndSplit_pdfFile(tmp2)
            new_docs.extend(load_file(docxPath,files.name))
        else:
            new_docs.extend(load_file(tmp,files.name))

files_uploaded = [f.name for f in path]if path else []
files_changed  = (files_uploaded!= st.session_state.uploaded_files)

if len(new_docs) > 0 and files_changed:
   # st.session_state.query = ""
    st.session_state.uploaded_files=files_uploaded
    st.session_state.docs = new_docs
    st.session_state.resume_verified = False   
    with st.spinner("Building dataset, please wait..."):       
        st.session_state.qa =  createRetrieverQA(new_docs,st.session_state.emb) 
        st.session_state.resumeUploaded = CheckIfResumeUploaded(st.session_state.qa)  
else:
    if not len(new_docs):
        st.info("Upload a resume or other supporting documents to begin")
        st.session_state.query=""
        st.session_state.firstPromptGiven=False
        st.stop()
        

     
if st.sidebar.button("Clear Chat"):
    #  st.text_area("Conversation with HRBuddy", value="", height=400)
    st.session_state.clear_chat_clicked = True
    st.session_state.singleConv = []                     
else:          
    st.session_state.clear_chat_clicked = False
              
if st.session_state.qa: 
    start_chat(st.session_state.qa,st.session_state.resumeUploaded)

textBytes = bytes()
if  st.session_state.singleConv != []:
    for q,r in st.session_state.singleConv:
        chat_line = f"You: {q}\nHRBuddy: {r}\n{'_' * 80}\n\n"
        textBytes += chat_line.encode('utf-8')
    
    now1 = datetime.now()
    st.sidebar.download_button(label = "Download Chat",
    data      = textBytes,
    file_name = "downloadedChat_"+ now1.strftime("%Y-%m-%d-%H:%M:%S")+".txt",
    mime      = "text/plain"
    )
# st.session_state.singleConv = []   

if not path:
     st.info("Upload a resume or other supporting documents to begin")
